"""AI 분석 결과 → Word(.docx) 다운로드 빌더.

analysis_results.raw_analysis JSONB(분석 LLM의 완전한 출력)를 문서 4 섹션으로 변환:
- 사업 개요  (key-value 2열 표)
- 평가항목   (4열 표)
- 요구사항   (그룹별 4열 표)
- 독소조항   (종합 위험도 + 6열 표)

ISP/ISMP·범용 둘 다 동일 스키마이므로 분기 없음.
"""
from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


_FONT_NAME = "Malgun Gothic"
_HEADER_FILL = "D9E1F2"
_LABEL_FILL = "F2F2F2"
_RISK_COLOR = {
    "danger": "C00000",
    "warning": "ED7D31",
    "caution": "BF8F00",
    "safe": "548235",
}
_SEVERITY_LABEL = {"danger": "위험", "warning": "경고", "caution": "주의", "safe": "안전"}


def _set_cell_shading(cell, fill_hex: str) -> None:
    """셀 배경색 (HEX, 예 'D9E1F2')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_run_font(run, *, size: int = 10, bold: bool = False, color: str | None = None) -> None:
    run.font.name = _FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # 한글 폰트 명시 (Word 호환)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), _FONT_NAME)
    r_fonts.set(qn("w:ascii"), _FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), _FONT_NAME)


def _add_paragraph(doc, text: str, *, size: int = 10, bold: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)


def _set_cell_text(cell, text: str, *, size: int = 10, bold: bool = False, color: str | None = None) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 기존 paragraph 재사용 (빈 셀은 paragraph 1개 가짐)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 기존 run 비우고 새로 추가
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text or "")
    _set_run_font(run, size=size, bold=bold, color=color)


def _format_won(val: Any) -> str:
    if val is None or val == "":
        return ""
    try:
        n = int(val)
        return f"{n:,} 원"
    except (ValueError, TypeError):
        return str(val)


# ---------------------------------------------------------------------------
# 섹션 빌더
# ---------------------------------------------------------------------------


def _add_section_heading(doc, text: str) -> None:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    _set_run_font(run, size=14, bold=True)


def _build_basic_info_section(doc, basic: dict, qualification: str, tech_requirements: list) -> None:
    rows: list[tuple[str, str]] = [
        ("발주기관", basic.get("issuing_org") or ""),
        ("사업 유형", basic.get("project_type") or ""),
        ("사업 개요", basic.get("project_summary") or ""),
        ("사업 범위", basic.get("project_scope") or ""),
        ("사업기간", basic.get("project_duration") or ""),
        ("추정가격(부가세 별도)", _format_won(basic.get("estimated_price"))),
        ("배정예산(부가세 포함)", _format_won(basic.get("allocated_budget"))),
        ("입찰 마감일시", basic.get("deadline") or ""),
        ("계약방법", basic.get("contract_method") or ""),
        ("참가 자격", qualification or ""),
        ("기술 요구사항", ", ".join(tech_requirements) if tech_requirements else ""),
    ]
    _add_section_heading(doc, "1. 사업 개요")
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(13)
    for i, (label, value) in enumerate(rows):
        lc, vc = table.rows[i].cells
        _set_cell_text(lc, label, bold=True)
        _set_cell_shading(lc, _LABEL_FILL)
        _set_cell_text(vc, value)


def _build_eval_section(doc, criteria: list) -> None:
    _add_section_heading(doc, "2. 평가항목")
    if not criteria:
        _add_paragraph(doc, "(평가항목 정보 없음)", size=10)
        return
    headers = ["평가부문", "세부 항목", "배점", "평가방법"]
    table = doc.add_table(rows=len(criteria) + 1, cols=len(headers))
    table.style = "Table Grid"
    for col, label in enumerate(headers):
        c = table.rows[0].cells[col]
        _set_cell_text(c, label, bold=True)
        _set_cell_shading(c, _HEADER_FILL)
    for i, item in enumerate(criteria, start=1):
        if not isinstance(item, dict):
            continue
        for col, key in enumerate(("category", "item", "score", "eval_method")):
            v = item.get(key)
            _set_cell_text(table.rows[i].cells[col], "" if v is None else str(v))


def _build_requirements_section(doc, requirements: dict) -> None:
    _add_section_heading(doc, "3. 요구사항")
    groups = requirements.get("groups") if isinstance(requirements, dict) else None
    if not groups:
        _add_paragraph(doc, "(요구사항 정보 없음)", size=10)
        return
    for group in groups:
        gname = (group.get("group_name") or "요구사항").strip()
        items = group.get("items") or []
        if not items:
            continue
        # 그룹 헤더
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(gname)
        _set_run_font(run, size=11, bold=True)
        # 테이블
        headers = ["ID", "요구사항명", "상세 내용"]
        table = doc.add_table(rows=len(items) + 1, cols=len(headers))
        table.style = "Table Grid"
        table.columns[0].width = Cm(2.5)
        table.columns[1].width = Cm(5)
        table.columns[2].width = Cm(10)
        for col, label in enumerate(headers):
            c = table.rows[0].cells[col]
            _set_cell_text(c, label, bold=True)
            _set_cell_shading(c, _HEADER_FILL)
        for i, it in enumerate(items, start=1):
            if not isinstance(it, dict):
                continue
            _set_cell_text(table.rows[i].cells[0], (it.get("id") or "").strip())
            _set_cell_text(table.rows[i].cells[1], (it.get("name") or "").strip())
            _set_cell_text(table.rows[i].cells[2], (it.get("description") or "").strip())


def _build_poison_section(doc, poison: dict) -> None:
    _add_section_heading(doc, "4. 독소조항")
    risk_level = (poison.get("risk_level") or "").lower() if isinstance(poison, dict) else ""
    summary = (poison.get("summary") or "").strip() if isinstance(poison, dict) else ""
    risk_label = _SEVERITY_LABEL.get(risk_level, risk_level or "—")
    risk_color = _RISK_COLOR.get(risk_level)

    # 종합 위험도 줄
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run("종합 위험도: ")
    _set_run_font(run1, size=11, bold=True)
    run2 = p.add_run(risk_label)
    _set_run_font(run2, size=11, bold=True, color=risk_color)
    if summary:
        _add_paragraph(doc, summary, size=10)

    items = poison.get("items") if isinstance(poison, dict) else None
    if not items:
        _add_paragraph(doc, "(독소조항 없음 또는 분석 결과에 항목 없음)", size=10)
        return

    headers = ["체크 ID", "위험도", "독소조항 원문", "판단 근거", "페이지", "출처"]
    table = doc.add_table(rows=len(items) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.columns[0].width = Cm(1.4)
    table.columns[1].width = Cm(1.4)
    table.columns[2].width = Cm(6.0)
    table.columns[3].width = Cm(4.5)
    table.columns[4].width = Cm(1.0)
    table.columns[5].width = Cm(3.0)
    for col, label in enumerate(headers):
        c = table.rows[0].cells[col]
        _set_cell_text(c, label, bold=True)
        _set_cell_shading(c, _HEADER_FILL)
    for i, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        sev = (item.get("severity") or "").lower()
        sev_label = _SEVERITY_LABEL.get(sev, sev or "")
        sev_color = _RISK_COLOR.get(sev)
        page_val = item.get("page")
        page_str = "" if page_val in (None, "") else str(page_val)
        _set_cell_text(table.rows[i].cells[0], (item.get("category") or "").strip(), bold=True)
        _set_cell_text(table.rows[i].cells[1], sev_label, bold=True, color=sev_color)
        _set_cell_text(table.rows[i].cells[2], (item.get("clause") or "").strip())
        _set_cell_text(table.rows[i].cells[3], (item.get("reason") or "").strip())
        _set_cell_text(table.rows[i].cells[4], page_str)
        _set_cell_text(table.rows[i].cells[5], (item.get("source") or "").strip())


# ---------------------------------------------------------------------------
# 공개 API
# ---------------------------------------------------------------------------


def build_analysis_docx(raw_analysis: dict, *, project_name: str | None = None) -> bytes:
    """analysis_results.raw_analysis → Word(.docx) 바이트."""
    if not isinstance(raw_analysis, dict):
        raw_analysis = {}

    basic = raw_analysis.get("basic_info") or {}
    eval_criteria = raw_analysis.get("eval_criteria") or []
    requirements = raw_analysis.get("requirements") or {}
    poison = raw_analysis.get("poison_clauses") or {}
    qualification = raw_analysis.get("qualification") or ""
    tech_requirements = raw_analysis.get("tech_requirements") or []

    doc = Document()
    # 한국어 호환 — 본문 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = _FONT_NAME
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), _FONT_NAME)

    # 표지
    title_text = project_name or basic.get("issuing_org") or "AI 분석 결과"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{title_text}\nAI 분석 결과")
    _set_run_font(run, size=18, bold=True)
    p.paragraph_format.space_after = Pt(18)

    _build_basic_info_section(doc, basic, qualification, tech_requirements)
    _build_eval_section(doc, eval_criteria)
    _build_requirements_section(doc, requirements)
    _build_poison_section(doc, poison)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
